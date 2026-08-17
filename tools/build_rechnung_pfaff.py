from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

try:
    from PIL import Image, ImageDraw, ImageFont
except Exception as exc:
    raise SystemExit(f"Pillow is required for logo generation: {exc}")


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "documents"
OUT_DIR.mkdir(exist_ok=True)

DOCX_PATH = OUT_DIR / "Rechnung_Rund_um_den_Baum_2000_EUR.docx"
LOGO_PATH = OUT_DIR / "miry-invoice-logo.png"

NAVY = RGBColor(28, 45, 53)
INK = RGBColor(42, 45, 43)
MUTED = RGBColor(102, 112, 108)
GOLD = RGBColor(199, 163, 92)
CREAM = "F7F2E8"
SOFT_GREEN = "E8EFE7"
BORDER = "D8C9AB"


def euro(value):
    return f"{value:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".") + " EUR"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="8"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=120, start=140, bottom=120, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    total_width = sum(int(width.twips) for width in widths)
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_width))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    old_grid = table._tbl.tblGrid
    if old_grid is not None:
        table._tbl.remove(old_grid)
    tbl_grid = OxmlElement("w:tblGrid")
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width.twips)))
        tbl_grid.append(grid_col)
    table._tbl.insert(1, tbl_grid)

    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width.twips)))
            tc_w.set(qn("w:type"), "dxa")
            row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(row.cells[idx])


def add_run(paragraph, text, size=10.5, color=INK, bold=False, italic=False, font="Arial"):
    run = paragraph.add_run(text)
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    return run


def add_label_value(cell, label, value):
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    add_run(p, label.upper(), size=7.5, color=MUTED, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    add_run(p2, value, size=10.5, color=INK, bold=True)


def make_logo():
    img = Image.new("RGBA", (760, 280), (255, 255, 255, 0))
    draw = ImageDraw.Draw(img)
    # soft layered mark inspired by clean/wild/nude/violet-style palettes
    draw.rounded_rectangle((20, 30, 250, 250), radius=64, fill=(247, 242, 232, 255), outline=(199, 163, 92, 255), width=5)
    draw.ellipse((80, 12, 220, 152), fill=(232, 239, 231, 230))
    draw.pieslice((44, 76, 226, 258), start=205, end=358, fill=(60, 82, 65, 180))
    draw.arc((68, 60, 226, 222), start=205, end=340, fill=(107, 84, 136, 190), width=7)
    draw.line((118, 176, 158, 95, 204, 178), fill=(28, 45, 53, 255), width=9, joint="curve")
    draw.line((154, 96, 154, 198), fill=(199, 163, 92, 255), width=5)

    def font(size, bold=False):
        candidates = [
            "C:/Windows/Fonts/playfairdisplay-bold.ttf",
            "C:/Windows/Fonts/georgiab.ttf" if bold else "C:/Windows/Fonts/georgia.ttf",
            "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        ]
        for candidate in candidates:
            try:
                return ImageFont.truetype(candidate, size)
            except Exception:
                pass
        return ImageFont.load_default()

    draw.text((290, 70), "MIRY", fill=(28, 45, 53, 255), font=font(70, True))
    draw.text((294, 148), "WEB DESIGN & DEVELOPMENT", fill=(102, 112, 108, 255), font=font(22, False))
    draw.line((294, 188, 590, 188), fill=(199, 163, 92, 255), width=3)
    img.save(LOGO_PATH)


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    add_run(p, text, size=13, color=NAVY, bold=True)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    add_run(p, text, size=9.5, color=INK)


def build_doc():
    make_logo()
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.top_margin = Cm(1.55)
    section.bottom_margin = Cm(1.55)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(header, "Rechnung | Webentwicklung", size=8.5, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(footer, "MIRY Web Design & Development - vertraulich - Seite ", size=8, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer._p.append(fld)

    top = doc.add_table(rows=1, cols=2)
    set_table_width(top, [Inches(3.15), Inches(2.85)])
    top.rows[0].cells[0].paragraphs[0].add_run().add_picture(str(LOGO_PATH), width=Inches(2.7))
    right = top.rows[0].cells[1]
    right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(right.paragraphs[0], "RECHNUNG", size=25, color=NAVY, bold=True)
    p = right.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p, "Nr. RE-2026-001\nDatum: 17.08.2026\nLeistungszeitraum: Juli-August 2026", size=9.5, color=MUTED)
    for cell in top.rows[0].cells:
        set_cell_border(cell, "FFFFFF", "0")

    doc.add_paragraph()
    parties = doc.add_table(rows=1, cols=2)
    set_table_width(parties, [Inches(3.0), Inches(3.0)])
    labels = [
        ("Rechnung von", "[Dein Name / Deine Firma]\n[Adresse]\n[PLZ Ort]\n[Steuernummer / USt-ID]\n[E-Mail]\n[Telefon]"),
        ("Rechnung an", "Joshua Pfaff\nRund um den Baum - Baumpflege\nMöllekensfeld 10\n46569 Hünxe"),
    ]
    for idx, (label, value) in enumerate(labels):
        cell = parties.rows[0].cells[idx]
        set_cell_shading(cell, CREAM if idx == 0 else SOFT_GREEN)
        set_cell_border(cell)
        add_label_value(cell, label, value)

    add_heading(doc, "Projekt")
    p = doc.add_paragraph()
    add_run(p, "Individuelle Website-Entwicklung für Rund um den Baum - Fa. Pfaff", size=11.5, color=INK, bold=True)
    p2 = doc.add_paragraph()
    add_run(
        p2,
        "Umsetzung einer mehrseitigen, responsiven Unternehmenswebsite inklusive Design-Feinschliff, "
        "Custom Code, Medienintegration, Kontaktformular-Vorbereitung und technischer Grundstruktur.",
        size=9.8,
        color=MUTED,
    )

    add_heading(doc, "Leistungsübersicht")
    for item in [
        "HTML5-Seitenstruktur, Inhaltsbereiche, wiederkehrende Komponenten und interne Verlinkungen",
        "Individuelles CSS-Design: Typografie, Farben, Abstände, Navigation, Karten, Sektionen und Hover-Zustände",
        "JavaScript-Interaktionen, mobile Navigation, UI-Verhalten, Animationen und visuelle Effekte",
        "Medienintegration inklusive HD-Video-Hero, Bildmaterial, generierte Bildbereiche und Optimierung der Darstellung",
        "Einbindung und Gestaltung der Google-Bewertungen im bestehenden Website-Design",
        "Kontaktformular mit Feldern, Benutzerführung, Foto-Upload-Vorbereitung und Danke-Seite",
        "Responsive Entwicklung für Smartphone, Tablet und Desktop inklusive Korrektur von Darstellungsproblemen",
        "Sonderseiten und technische Details: 404-Seite, Danke-Seite, Favicon, OpenGraph-Metadaten, SEO-Grundlagen, Debugging und Testing",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Noch offene finale Einrichtung")
    for item in [
        "Domain- und Hosting-Setup inklusive HTTPS/SSL-Konfiguration",
        "Finale E-Mail-Anbindung des Kontaktformulars über das Hosting",
        "Optional: CMS/Admin-Zugang zur späteren eigenständigen Pflege von Fotos und Texten",
    ]:
        add_bullet(doc, item)

    doc.add_page_break()
    add_heading(doc, "Abrechnung")
    table = doc.add_table(rows=1, cols=4)
    set_table_width(table, [Inches(2.85), Inches(0.7), Inches(1.15), Inches(1.3)])
    headers = ["Leistung", "Anz.", "Einzelpreis", "Betrag"]
    for idx, label in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, "1C2D35")
        set_cell_border(cell, "1C2D35")
        p = cell.paragraphs[0]
        add_run(p, label, size=8.5, color=RGBColor(255, 255, 255), bold=True)
    row = table.add_row()
    values = [
        "Website-Entwicklung Rund um den Baum gemäß Leistungsübersicht",
        "1",
        euro(2000),
        euro(2000),
    ]
    for idx, value in enumerate(values):
        cell = row.cells[idx]
        set_cell_border(cell)
        if idx == 0:
            add_run(cell.paragraphs[0], value, size=9.2, color=INK)
        else:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            add_run(cell.paragraphs[0], value, size=9.2, color=INK, bold=idx == 3)

    total = doc.add_table(rows=3, cols=2)
    set_table_width(total, [Inches(4.15), Inches(1.85)])
    total_data = [
        ("Zwischensumme", euro(2000)),
        ("Umsatzsteuer", "[bitte ergänzen]"),
        ("Rechnungsbetrag", euro(2000)),
    ]
    for r, (label, value) in enumerate(total_data):
        for c in range(2):
            cell = total.rows[r].cells[c]
            set_cell_border(cell, "FFFFFF", "0")
            if r == 2:
                set_cell_shading(cell, CREAM)
        total.rows[r].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        total.rows[r].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_run(total.rows[r].cells[0].paragraphs[0], label, size=9.8, color=NAVY if r == 2 else MUTED, bold=r == 2)
        add_run(total.rows[r].cells[1].paragraphs[0], value, size=10.5, color=NAVY if r == 2 else INK, bold=r == 2)

    add_heading(doc, "Zahlungsinformationen")
    info = doc.add_table(rows=1, cols=2)
    set_table_width(info, [Inches(3.0), Inches(3.0)])
    for idx, (label, value) in enumerate([
        ("Zahlungsziel", "zahlbar innerhalb von 7 Tagen nach Rechnungserhalt"),
        ("Bankverbindung", "IBAN: [bitte ergänzen]\nBIC: [bitte ergänzen]\nKontoinhaber: [bitte ergänzen]"),
    ]):
        cell = info.rows[0].cells[idx]
        set_cell_shading(cell, SOFT_GREEN if idx == 0 else CREAM)
        set_cell_border(cell)
        add_label_value(cell, label, value)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    add_run(
        p,
        "Hinweis: Bitte Steuernummer/USt-ID, vollständige Rechnungsadresse, Bankverbindung und den passenden Umsatzsteuerhinweis vor Versand prüfen und ergänzen.",
        size=8.5,
        color=MUTED,
        italic=True,
    )

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
    print(DOCX_PATH.name)
