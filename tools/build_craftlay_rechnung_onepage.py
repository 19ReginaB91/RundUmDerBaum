from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "documents"
OUT_DIR.mkdir(exist_ok=True)

DOCX_PATH = OUT_DIR / "Rechnung_CraftLay_Rund_um_den_Baum_2026-1001_onepage.docx"
LOGO_SOURCE = Path(
    r"C:\Users\Admin\.codex\generated_images\019f9539-28b7-7c40-954f-e0766e1ae4f5\call_mVZYPKd1JWUn9rAsV6ZCtA2L.png"
)
LOGO_PATH = OUT_DIR / "craftlay-cl-logo.png"

if LOGO_SOURCE.exists():
    LOGO_PATH.write_bytes(LOGO_SOURCE.read_bytes())


INK = RGBColor(31, 37, 38)
MUTED = RGBColor(95, 100, 97)
VIOLET = RGBColor(88, 35, 118)
GOLD = RGBColor(185, 132, 43)
WHITE = RGBColor(255, 255, 255)
LINE = "D6BF96"
DARK = "1F2526"
NUDE = "F5EEE4"
VIOLET_FILL = "F0E7F4"
TOTAL_FILL = "FBF6EE"


def euro(value):
    return f"{value:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".") + " EUR"


def run(p, text, font="Arial", size=9.5, color=INK, bold=False, italic=False):
    r = p.add_run(text)
    r.font.name = font
    r._element.rPr.rFonts.set(qn("w:ascii"), font)
    r._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.bold = bold
    r.italic = italic
    return r


def shade(cell, fill):
    pr = cell._tc.get_or_add_tcPr()
    node = pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        pr.append(node)
    node.set(qn("w:fill"), fill)


def border(cell, color=LINE, size="8"):
    pr = cell._tc.get_or_add_tcPr()
    borders = pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def margins(cell, top=95, start=120, bottom=95, end=120):
    pr = cell._tc.get_or_add_tcPr()
    mar = pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        pr.append(mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def fixed(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    total = sum(int(w.twips) for w in widths)
    pr = table._tbl.tblPr
    tw = pr.find(qn("w:tblW"))
    if tw is None:
        tw = OxmlElement("w:tblW")
        pr.append(tw)
    tw.set(qn("w:w"), str(total))
    tw.set(qn("w:type"), "dxa")
    layout = pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    old = table._tbl.tblGrid
    if old is not None:
        table._tbl.remove(old)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width.twips)))
        grid.append(col)
    table._tbl.insert(1, grid)

    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = width
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            pr = cell._tc.get_or_add_tcPr()
            cw = pr.find(qn("w:tcW"))
            if cw is None:
                cw = OxmlElement("w:tcW")
                pr.append(cw)
            cw.set(qn("w:w"), str(int(width.twips)))
            cw.set(qn("w:type"), "dxa")
            margins(cell)


def label_cell(cell, label, value, fill):
    shade(cell, fill)
    border(cell)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    run(p, label.upper(), size=7.2, color=MUTED, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    run(p2, value, size=9.2, bold=True)


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(1.2)
    sec.bottom_margin = Cm(1.2)
    sec.left_margin = Cm(1.25)
    sec.right_margin = Cm(1.25)
    sec.header_distance = Cm(0.55)
    sec.footer_distance = Cm(0.55)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(9.5)

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(
        footer,
        "CraftLay | Regina Bolotnikova | regina.craftlay@gmail.com",
        size=7.7,
        color=MUTED,
    )

    top = doc.add_table(rows=1, cols=2)
    fixed(top, [Inches(3.2), Inches(3.55)])
    for cell in top.rows[0].cells:
        border(cell, "FFFFFF", "0")
    if LOGO_PATH.exists():
        top.rows[0].cells[0].paragraphs[0].add_run().add_picture(str(LOGO_PATH), width=Inches(1.55))
    p = top.rows[0].cells[1].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run(p, "RECHNUNG", font="Georgia", size=24, color=VIOLET, bold=True)
    p2 = top.rows[0].cells[1].add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run(p2, "Rechnung Nr. 2026-1001", size=9.5, color=INK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run(p, "Regina Bolotnikova, Inhaberin", size=8.5, color=VIOLET, bold=True)

    info = doc.add_table(rows=1, cols=2)
    fixed(info, [Inches(3.25), Inches(3.5)])
    left, right = info.rows[0].cells
    label_cell(
        left,
        "Rechnung an",
        "Joshua Pfaff\nRund um den Baum - Baumpflege\nMöllekensfeld 10\n46569 Hünxe\nuli.p.pfaff@t-online.de",
        NUDE,
    )
    label_cell(
        right,
        "Rechnungsdaten",
        "Rechnungsdatum: 17.08.2026\nLeistungszeitraum: Juli-August 2026\nStunden: 172 Std.\nFällig bis: 31.08.2026",
        VIOLET_FILL,
    )

    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(4)
    run(h, "Leistungsübersicht", font="Georgia", size=13, color=VIOLET, bold=True)

    table = doc.add_table(rows=1, cols=5)
    fixed(table, [Inches(2.55), Inches(1.2), Inches(0.75), Inches(1.05), Inches(1.2)])
    headers = ["Beschreibung", "Zeitraum", "Menge", "Einzelpreis", "Gesamt"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade(cell, DARK)
        border(cell, DARK)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
        run(p, text, size=8.2, color=WHITE, bold=True)

    row = table.add_row()
    values = [
        "Webdesign & Webentwicklung Rund um den Baum",
        "Juli-August 2026",
        "172 Std.",
        "pauschal",
        euro(2000),
    ]
    for idx, value in enumerate(values):
        cell = row.cells[idx]
        border(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if idx in (2, 3, 4) else WD_ALIGN_PARAGRAPH.LEFT
        run(p, value, size=8.9, bold=idx == 4)

    totals = doc.add_table(rows=3, cols=2)
    fixed(totals, [Inches(5.3), Inches(1.45)])
    for r_idx, (label, value) in enumerate(
        [
            ("Zwischensumme", euro(2000)),
            ("Umsatzsteuer", "[bitte ergänzen]"),
            ("Rechnungsbetrag", euro(2000)),
        ]
    ):
        for c_idx in range(2):
            cell = totals.rows[r_idx].cells[c_idx]
            border(cell, "FFFFFF", "0")
            if r_idx == 2:
                shade(cell, TOTAL_FILL)
        totals.rows[r_idx].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        totals.rows[r_idx].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run(totals.rows[r_idx].cells[0].paragraphs[0], label, size=8.8, color=VIOLET if r_idx == 2 else MUTED, bold=r_idx == 2)
        run(totals.rows[r_idx].cells[1].paragraphs[0], value, size=9.4, bold=r_idx == 2)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(5)
    note.paragraph_format.space_after = Pt(5)
    run(
        note,
        "Hinweis: Bitte Umsatzsteuerhinweis/Steuernummer vor dem Versand ergänzen.",
        size=7.8,
        color=MUTED,
        italic=True,
    )

    bank = doc.add_table(rows=2, cols=3)
    fixed(bank, [Inches(2.25), Inches(2.25), Inches(2.25)])
    entries = [
        ("Zahlungsempfänger", "Regina Bolotnikova"),
        ("IBAN", "[bitte eintragen]"),
        ("BIC / Bank", "[bitte eintragen]"),
        ("Adresse", "Reicher-Leute-Stege 12\n46485 Wesel"),
        ("Steuernummer", "[bitte eintragen]"),
        ("Kontakt", "regina.craftlay@gmail.com"),
    ]
    for idx, (label, value) in enumerate(entries):
        cell = bank.rows[idx // 3].cells[idx % 3]
        label_cell(cell, label, value, NUDE if idx % 2 == 0 else VIOLET_FILL)

    thanks = doc.add_paragraph()
    thanks.paragraph_format.space_before = Pt(10)
    run(thanks, "Vielen Dank für die Zusammenarbeit.", font="Georgia", size=11.2, color=VIOLET, bold=True)

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build()
    print(DOCX_PATH.name)
