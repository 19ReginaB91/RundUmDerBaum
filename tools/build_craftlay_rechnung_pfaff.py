from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "documents"
OUT_DIR.mkdir(exist_ok=True)

DOCX_PATH = OUT_DIR / "Rechnung_CraftLay_Rund_um_den_Baum_2026-1001.docx"
LOGO_SOURCE = Path(
    r"C:\Users\Admin\.codex\generated_images\019f9539-28b7-7c40-954f-e0766e1ae4f5\call_mVZYPKd1JWUn9rAsV6ZCtA2L.png"
)
LOGO_PATH = OUT_DIR / "craftlay-cl-logo.png"

if LOGO_SOURCE.exists() and not LOGO_PATH.exists():
    LOGO_PATH.write_bytes(LOGO_SOURCE.read_bytes())


INK = RGBColor(31, 37, 38)
MUTED = RGBColor(98, 105, 101)
VIOLET = RGBColor(83, 32, 111)
GOLD = RGBColor(184, 132, 43)
PAPER = "FBF7EF"
NUDE = "F2E8DA"
SOFT_VIOLET = "EEE6F2"
LINE = "D9C7A7"
DARK = "1F2526"


def euro(value):
    return f"{value:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".") + " EUR"


def set_run(run, font="Georgia", size=10, color=INK, bold=False, italic=False):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def add_text(paragraph, text, font="Georgia", size=10, color=INK, bold=False, italic=False):
    run = paragraph.add_run(text)
    set_run(run, font, size, color, bold, italic)
    return run


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def borders(cell, color=LINE, size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        node = tc_borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def cell_margins(cell, top=120, start=150, bottom=120, end=150):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def fixed_table(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    total = sum(int(w.twips) for w in widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    old_grid = table._tbl.tblGrid
    if old_grid is not None:
        table._tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width.twips)))
        grid.append(col)
    table._tbl.insert(1, grid)

    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = width
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width.twips)))
            tc_w.set(qn("w:type"), "dxa")
            cell_margins(cell)


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(13)
    p.paragraph_format.space_after = Pt(5)
    add_text(p, text, size=13, color=VIOLET, bold=True)
    return p


def label_value(cell, label, value):
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    add_text(p, label.upper(), font="Arial", size=7.5, color=MUTED, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    add_text(p2, value, font="Arial", size=9.7, color=INK, bold=True)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.55)
    p.paragraph_format.space_after = Pt(3)
    add_text(p, text, font="Arial", size=9.6, color=INK)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.35)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.55)
    section.right_margin = Cm(1.55)
    section.header_distance = Cm(0.75)
    section.footer_distance = Cm(0.75)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(header, "CraftLay | Rechnung 2026-1001", font="Arial", size=8.5, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(footer, "CraftLay - Regina Bolotnikova - Design & Development", font="Arial", size=8, color=MUTED)

    top = doc.add_table(rows=1, cols=2)
    fixed_table(top, [Inches(2.25), Inches(4.0)])
    for c in top.rows[0].cells:
        borders(c, "FFFFFF", "0")

    logo_cell, title_cell = top.rows[0].cells
    if LOGO_PATH.exists():
        logo_cell.paragraphs[0].add_run().add_picture(str(LOGO_PATH), width=Inches(1.55))

    title_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(title_cell.paragraphs[0], "RECHNUNG", font="Georgia", size=27, color=INK, bold=True)
    p = title_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(p, "Nr. 2026-1001\nDatum: 17.08.2026\nLeistungszeitraum: 172 Stunden", font="Arial", size=9.5, color=MUTED)

    accent = doc.add_table(rows=1, cols=3)
    fixed_table(accent, [Inches(2.7), Inches(0.85), Inches(2.7)])
    for idx, cell in enumerate(accent.rows[0].cells):
        borders(cell, "FFFFFF", "0")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if idx == 1:
            add_text(p, "•  •  •", font="Arial", size=14, color=GOLD, bold=True)
        else:
            add_text(p, "────────────", font="Arial", size=8, color=GOLD)

    doc.add_paragraph()
    parties = doc.add_table(rows=1, cols=2)
    fixed_table(parties, [Inches(3.1), Inches(3.1)])
    data = [
        (
            "Rechnung von",
            "Regina Bolotnikova\nCraftLay\nReicher-Leute-Stege 12\n46485 Wesel\nregina.craftlay@gmail.com\nSteuernummer: [bitte ergänzen]",
            NUDE,
        ),
        (
            "Rechnung an",
            "Joshua Pfaff\nRund um den Baum - Baumpflege\nMöllekensfeld 10\n46569 Hünxe\nuli.p.pfaff@t-online.de",
            SOFT_VIOLET,
        ),
    ]
    for idx, (label, value, fill) in enumerate(data):
        cell = parties.rows[0].cells[idx]
        shade(cell, fill)
        borders(cell, LINE)
        label_value(cell, label, value)

    heading(doc, "Projektumfang")
    p = doc.add_paragraph()
    add_text(
        p,
        "Individuelle Website-Entwicklung für Rund um den Baum - Fa. Pfaff.",
        font="Arial",
        size=10.5,
        color=INK,
        bold=True,
    )
    p2 = doc.add_paragraph()
    add_text(
        p2,
        "Erstellung und Ausarbeitung eines responsiven Unternehmensauftritts mit individuellem Design, "
        "Custom Code, Seitenstruktur, Medienintegration und vorbereiteter Formularlogik.",
        font="Arial",
        size=9.6,
        color=MUTED,
    )

    heading(doc, "Kurz zusammengefasst wurde umgesetzt")
    for item in [
        "Mehrseitige HTML/CSS/JavaScript-Website mit sauberer Seitenstruktur und interner Navigation",
        "Individuelles responsives Design für Desktop, Tablet und Smartphone",
        "Startseite, Leistungsseiten, Über-uns-Bereich, Arbeiten/Galerie, Kontakt, Impressum, Datenschutz, 404- und Danke-Seite",
        "Integration und Optimierung von Bildern, Video-Hero, generierten Arbeits- und Leistungsbildern",
        "Kontaktformular mit Foto-Upload-Vorbereitung, Danke-Seite und späterer Netlify-Forms-Anbindung",
        "Google-Bewertungsbereich, mobile Navigation, Favicon, OpenGraph-Grundlagen, technische Tests und Feinschliff",
    ]:
        bullet(doc, item)

    doc.add_page_break()
    heading(doc, "Abrechnung")
    table = doc.add_table(rows=1, cols=4)
    fixed_table(table, [Inches(3.0), Inches(0.75), Inches(1.15), Inches(1.3)])
    headers = ["Leistung", "Std.", "Satz", "Betrag"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade(cell, DARK)
        borders(cell, DARK)
        p = cell.paragraphs[0]
        add_text(p, text, font="Arial", size=8.8, color=RGBColor(255, 255, 255), bold=True)

    row = table.add_row()
    values = [
        "Pauschale Website-Entwicklung gemäß Projektumfang",
        "172",
        "pauschal",
        euro(2000),
    ]
    for idx, value in enumerate(values):
        cell = row.cells[idx]
        borders(cell, LINE)
        if idx == 0:
            add_text(cell.paragraphs[0], value, font="Arial", size=9.3, color=INK)
        else:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            add_text(cell.paragraphs[0], value, font="Arial", size=9.3, color=INK, bold=idx == 3)

    total = doc.add_table(rows=3, cols=2)
    fixed_table(total, [Inches(4.35), Inches(1.85)])
    lines = [
        ("Zwischensumme", euro(2000)),
        ("Umsatzsteuer", "[bitte ergänzen]"),
        ("Rechnungsbetrag", euro(2000)),
    ]
    for row_idx, (label, value) in enumerate(lines):
        for col_idx in range(2):
            cell = total.rows[row_idx].cells[col_idx]
            borders(cell, "FFFFFF", "0")
            if row_idx == 2:
                shade(cell, PAPER)
        total.rows[row_idx].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        total.rows[row_idx].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_text(total.rows[row_idx].cells[0].paragraphs[0], label, font="Arial", size=9.7, color=MUTED if row_idx < 2 else VIOLET, bold=row_idx == 2)
        add_text(total.rows[row_idx].cells[1].paragraphs[0], value, font="Arial", size=10.5, color=INK, bold=row_idx == 2)

    heading(doc, "Zahlungsinformationen")
    pay = doc.add_table(rows=1, cols=2)
    fixed_table(pay, [Inches(3.1), Inches(3.1)])
    for idx, (label, value) in enumerate([
        ("Zahlungsziel", "14 Tage nach Rechnungserhalt"),
        ("Bankverbindung", "IBAN: [bitte eintragen]\nBIC: [bitte eintragen]\nKontoinhaber: Regina Bolotnikova"),
    ]):
        cell = pay.rows[0].cells[idx]
        shade(cell, NUDE if idx == 0 else SOFT_VIOLET)
        borders(cell, LINE)
        label_value(cell, label, value, )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    add_text(
        p,
        "Vielen Dank für das Vertrauen und die angenehme Zusammenarbeit.",
        font="Georgia",
        size=11.5,
        color=VIOLET,
        italic=True,
    )

    note = doc.add_paragraph()
    add_text(
        note,
        "Hinweis: Bitte vor dem Versand IBAN/BIC, Steuernummer und den passenden Umsatzsteuerhinweis ergänzen.",
        font="Arial",
        size=8.6,
        color=MUTED,
        italic=True,
    )

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build()
    print(DOCX_PATH.name)
